#!/usr/bin/env python3
"""生成 task2 数据炼金术 Word 文档 (.docx)，可自行编辑"""

import os, re
from datetime import datetime

import numpy as np
import pandas as pd
import matplotlib.font_manager as fm
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

WORKDIR = os.path.dirname(os.path.abspath(__file__))
DATA_DIR = os.path.join(WORKDIR, "埃斯顿_002747")
CHART_PATH = os.path.join(DATA_DIR, "technical_chart.png")
OUTPUT = os.path.join(DATA_DIR, "report.docx")

df = pd.read_csv(os.path.join(DATA_DIR, "data.csv"), parse_dates=["date"])
closes = df["close"].values

# ---- 统计 ----
def calc_stats():
    s = {}
    for col in ["open","high","low","close","volume","amount"]:
        v = df[col].values.astype(float)
        s[col] = {"mean":np.mean(v),"median":np.median(v),"std":np.std(v),
                   "min":np.min(v),"max":np.max(v),"q1":np.percentile(v,25),"q3":np.percentile(v,75)}
    rets = np.diff(closes) / closes[:-1]
    s["return"] = {"mean":np.mean(rets),"std":np.std(rets),"max":np.max(rets),"min":np.min(rets),
                    "annual_ret":np.mean(rets)*252,"annual_vol":np.std(rets)*np.sqrt(252),
                    "sharpe":np.mean(rets)/np.std(rets)*np.sqrt(252) if np.std(rets)>0 else 0,
                    "up":int(np.sum(rets>0)),"down":int(np.sum(rets<0))}
    return s
S = calc_stats()
FIRST, LAST = df.iloc[0], df.iloc[-1]
TOTAL_CHANGE = (LAST["close"]-FIRST["close"])/FIRST["close"]*100

# 指标值
delta = df["close"].diff()
gain = delta.clip(lower=0); loss = (-delta).clip(lower=0)
ag = gain.ewm(alpha=1/14,min_periods=14,adjust=False).mean()
al = loss.ewm(alpha=1/14,min_periods=14,adjust=False).mean()
rsi_val = float((100-100/(1+ag/al)).iloc[-1])
e12=df["close"].ewm(span=12,adjust=False).mean(); e26=df["close"].ewm(span=26,adjust=False).mean()
dif_val = float((e12-e26).iloc[-1])
dea_s = (e12-e26).ewm(span=9,adjust=False).mean(); dea_val=float(dea_s.iloc[-1])
bbm=df["close"].rolling(20,min_periods=20).mean(); bbs=df["close"].rolling(20,min_periods=20).std(ddof=0)
bb_up=float((bbm+2*bbs).iloc[-1]); bb_md=float(bbm.iloc[-1]); bb_lo=float((bbm-2*bbs).iloc[-1])

# ===== 文档构建 =====
doc = Document()

# 页面设置
section = doc.sections[0]
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# 默认字体
style = doc.styles['Normal']
style.font.name = 'SimSun'
style.font.size = Pt(10.5)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
# 设置中文字体
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'SimSun'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(13)
    return h

def add_para(text, bold=False, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'SimSun'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(10.5)
    run.bold = bold
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if align:
        p.alignment = align
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def add_table(headers, rows):
    ncols = len(headers)
    table = doc.add_table(rows=len(rows)+1, cols=ncols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'SimSun'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run.font.size = Pt(9)
                run.bold = True
    # Body
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.name = 'SimSun'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.size = Pt(9)
    doc.add_paragraph()  # spacing after table

def add_caption(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'SimSun'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(9)
    run.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)

def add_image(path, width_inches=5.5):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width_inches))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def page_break():
    doc.add_page_break()

# ===== 封面 =====
for _ in range(6):
    doc.add_paragraph()
add_para("task2 数据炼金术：数据诊断与构造交易指标", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.paragraphs[-1].runs[0].font.size = Pt(22)
doc.add_paragraph()
add_para("量化分析课程作业", align=WD_ALIGN_PARAGRAPH.CENTER)
doc.paragraphs[-1].runs[0].font.size = Pt(14)
doc.add_paragraph()
add_para(f"分析标的：埃斯顿 (002747.SZ)", align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(f"数据区间：{FIRST['date'].strftime('%Y-%m-%d')} ~ {LAST['date'].strftime('%Y-%m-%d')}", align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(f"日期：{datetime.now().strftime('%Y年%m月%d日')}", align=WD_ALIGN_PARAGRAPH.CENTER)

page_break()

# ===== 一 =====
add_heading_styled("一、数据获取与诊断", level=1)

add_heading_styled("1.1 数据获取过程", level=2)
add_para("我通过 WorkBuddy AI 编程助手，调用 Tushare 金融数据接口（mcp__tushareMcp__daily），获取了埃斯顿（002747.SZ）的日 K 线数据。在对话中，我向 AI 提出了「通过 Tushare 获取埃斯顿近一年的交易价格和交易量信息，存储到本地」的需求。AI 自动完成了接口调用、日期格式化、成交额单位换算（Tushare 返回的金额单位为千元，需乘以 1000 转换为元），并将干净数据输出为 CSV 和 JSON 两种格式，存入 埃斯顿_002747/ 目录。")
add_para(f"最终获得 {len(df)} 个交易日的数据，时间跨度为 {FIRST['date'].strftime('%Y年%m月%d日')} 至 {LAST['date'].strftime('%Y年%m月%d日')}。价格从 {FIRST['close']:.2f} 元上涨至 {LAST['close']:.2f} 元，区间涨幅 {TOTAL_CHANGE:+.1f} %。")

add_heading_styled("1.2 数据质量检查", level=2)
add_para("数据质量是量化分析的基础。我在 AI 的协助下完成了以下检查：")
add_para("（1）缺失值扫描：对 open、high、low、close、volume、amount 六列逐一检查，结果为 0 条缺失，数据完整性 100 %。")
add_para("（2）日期断点分析：识别出 4 处间隔超过 4 个自然日的断点（2025年国庆 9 天、2026年元旦 5 天、2026年春节 11 天、2026年劳动节 6 天），经确认均为中国法定节假日休市，无数据丢失。")
add_para("（3）价格逻辑验证：逐条验证 low <= open <= high 且 low <= close <= high，全部通过，无异常极值。")

add_heading_styled("1.3 描述性统计", level=2)
add_para("以下为收盘价及日收益率的完整描述性统计，由 AI 代我执行 Node.js 计算得出。")
add_caption("表 1：收盘价描述性统计")
add_table(
    ["指标", "收盘价(元)", "成交量(手)", "成交额(万元)"],
    [
        ["均值", f"{S['close']['mean']:.2f}", f"{S['volume']['mean']/10000:.0f}", f"{S['amount']['mean']/10000:.0f}"],
        ["中位数", f"{S['close']['median']:.2f}", f"{S['volume']['median']/10000:.0f}", f"{S['amount']['median']/10000:.0f}"],
        ["标准差", f"{S['close']['std']:.2f}", f"{S['volume']['std']/10000:.0f}", f"{S['amount']['std']/10000:.0f}"],
        ["最小值", f"{S['close']['min']:.2f}", f"{S['volume']['min']/10000:.0f}", f"{S['amount']['min']/10000:.0f}"],
        ["最大值", f"{S['close']['max']:.2f}", f"{S['volume']['max']/10000:.0f}", f"{S['amount']['max']/10000:.0f}"],
        ["Q1 (25%)", f"{S['close']['q1']:.2f}", f"{S['volume']['q1']/10000:.0f}", f"{S['amount']['q1']/10000:.0f}"],
        ["Q3 (75%)", f"{S['close']['q3']:.2f}", f"{S['volume']['q3']/10000:.0f}", f"{S['amount']['q3']/10000:.0f}"],
    ]
)
add_para("从表 1 可以看出，收盘价右偏（均值 24.28 > 中位数 23.65），成交量分布同样右偏（均值 35 万手 > 中位数 24 万手），说明存在少数放量交易日拉高了整体均值。收盘价标准差 4.21 元，相对于 24.28 元的均值，变异系数约 17.3 %，波动性适中。")

add_caption("表 2：日收益率统计")
add_table(
    ["指标", "值"],
    [
        ["日均收益率", f"{S['return']['mean']*100:.4f}%"],
        ["标准差", f"{S['return']['std']*100:.4f}%"],
        ["最大单日涨幅", f"{S['return']['max']*100:.2f}%"],
        ["最大单日跌幅", f"{S['return']['min']*100:.2f}%"],
        ["年化收益率", f"{S['return']['annual_ret']*100:.2f}%"],
        ["年化波动率", f"{S['return']['annual_vol']*100:.2f}%"],
        ["夏普比率(0%)", f"{S['return']['sharpe']:.4f}"],
        ["上涨/下跌天数", f"{S['return']['up']} / {S['return']['down']} ({S['return']['up']/(S['return']['up']+S['return']['down'])*100:.1f}% / {S['return']['down']/(S['return']['up']+S['return']['down'])*100:.1f}%)"],
    ]
)
add_para("从表 2 可以看出，日收益率均值 +0.38 %，年化约 96.2 %，但标准差达 2.86 %（年化 45.4 %），表明波动较大。夏普比率 2.12，在 A 股中属于较高水平，但近一年正值机器人/智能制造赛道景气周期，该数值是否可持续需要警惕。上涨日与下跌日基本持平（120 vs 117），说明收益率中有正向偏态。")

# ===== 二 =====
page_break()
add_heading_styled("二、技术指标理论与计算方法", level=1)
add_para("在完成数据清洗后，我请 AI 解释了 RSI、MACD 和布林带的含义与计算方法。以下是整理后的理论内容。")

add_heading_styled("2.1 RSI — 相对强弱指标", level=2)
add_para("RSI（Relative Strength Index）由 J. Welles Wilder 于 1978 年提出，衡量价格变动的内部强度，取值范围 0 - 100。其核心思想是：将过去 N 日的平均涨幅与平均跌幅做比值，通过归一化得到反映多空力量对比的震荡指标。")
add_para("计算步骤：(1) 逐日计算涨跌：Gain[t] = max(Close[t]-Close[t-1], 0)，Loss[t] = max(Close[t-1]-Close[t], 0)；(2) 使用 Wilder 平滑法（alpha = 1/N）计算 N 日平均涨幅 AvgG 和平均跌幅 AvgL；(3) 计算 RS = AvgG/AvgL；(4) RSI = 100 - 100/(1+RS)。标准参数 N = 14。")
add_para("RSI > 70 通常视为超买，可能回调；RSI < 30 视为超卖，可能反弹；30 - 70 之间为正常波动区间。此外，RSI 的顶/底背离（价格新高而 RSI 未新高，或价格新低而 RSI 未新低）是重要的趋势反转信号。")

add_heading_styled("2.2 MACD — 指数平滑异同移动平均线", level=2)
add_para("MACD（Moving Average Convergence Divergence）由 Gerald Appel 于 1979 年提出，是经典的趋势跟踪兼动量指标。它通过短周期与长周期指数移动平均（EMA）的差值衡量趋势力度，再对该差值做平滑得到信号线。")
add_para("MACD 由三条曲线构成：DIF（快线）= EMA(12) - EMA(26)，反映短期与长期趋势的偏离；DEA（慢线/信号线）= EMA(DIF, 9)，是 DIF 的平滑线，减缓噪音；MACD 柱 = (DIF - DEA) x 2，正值表示多头动能，负值表示空头动能。")
add_para("信号规则：金叉（DIF 上穿 DEA）为多头入场信号，死叉（DIF 下穿 DEA）为空头出场信号。柱状线的增减速度同样具有参考价值。")

add_heading_styled("2.3 布林带 — Bollinger Bands", level=2)
add_para("布林带由 John Bollinger 于 1980 年代提出，基于统计学中正态分布的性质，用移动平均与标准差构建价格的动态波动区间。三条轨线为：中轨 = MA(Close, 20)，上轨 = MA(20) + 2 x sigma(20)，下轨 = MA(20) - 2 x sigma(20)，其中 sigma 为 20 日收盘价标准差。")
add_para("若价格服从正态分布，约 95 % 的观测值将落在上下轨之间。因此上轨可视为统计意义上的阻力参考，下轨为支撑参考。带宽收窄（Squeeze）预示波动率极低，通常伴随大幅突破；价格突破上轨表示短期强势但可能超买，跌破下轨表示短期弱势但可能超卖。")

# ===== 三 =====
page_break()
add_heading_styled("三、程序实现与图表解读", level=1)

add_heading_styled("3.1 编程实现过程", level=2)
add_para("在理解指标原理后，我请 AI 编写了一个通用的 Python 程序 technical_indicators.py。该程序使用 pandas 进行数据处理和指标计算，matplotlib 进行可视化，支持命令行参数调用。AI 在编写过程中经历了几轮调试：最初遇到了 K 线宽度计算不准确的问题（日历日 vs 交易日），以及中文字体渲染异常，最终通过使用整数索引替代日期坐标、自动检测系统字体得以解决。整个过程让我体会到，即使是标准的指标公式，落地到代码中也需要处理许多琐碎的工程细节。")

add_heading_styled("3.2 运行结果", level=2)
add_image(CHART_PATH, 5.5)
add_caption("图 1：埃斯顿(002747)技术指标综合分析 (2025.07 - 2026.07)")
add_para("图 1 为 technical_indicators.py 对埃斯顿 242 个交易日的输出结果，自上而下分别展示 K 线图（叠加布林带）、成交量、MACD 和 RSI。")

add_heading_styled("3.3 图表解读", level=2)
add_para(f"截至 2026 年 7 月 3 日收盘价 {LAST['close']:.2f} 元，各指标状态如下：")
add_para(f"布林带与 K 线（第一面板）：价格 {LAST['close']:.2f} 元已突破布林上轨 {bb_up:.2f} 元，属于统计意义上的强势突破。中轨（MA20）自 2026 年 3 月以来持续向上倾斜，短中期均线呈多头排列。")
add_para(f"MACD（第三面板）：DIF = {dif_val:.2f}，DEA = {dea_val:.2f}，DIF > DEA 维持金叉状态，柱线为正值。全区间共出现 7 次金叉和 7 次死叉，信号频率适中。当前 DIF 处于零轴上方且自 6 月初加速上行，表明多头动能尚未衰减。")
add_para(f"RSI（第四面板）：RSI(14) = {rsi_val:.2f}，已进入超买区间（> 70）。回顾区间历史，RSI 在 6 月初曾降至 21.12（超卖区），随后在一个月内迅速攀升至 70 以上，上升速度较快。")
add_para("综合来看，三指标均指向强势多头格局：布林带确认价格处于极端强势区，MACD 确认趋势方向向上且动能充沛，RSI 提示短期已经过热。这种指标共振既可能意味着趋势延续，也可能预示回调的临近。对于量化策略而言，此时更需要关注后续 RSI 回落至 70 以下和价格回归布林带内的反转信号。")

# ===== 四 =====
page_break()
add_heading_styled("四、工具部署与资源链接", level=1)

add_heading_styled("4.1 tushare-stock-loader 技能", level=2)
add_para("为便于日后复用，我请 AI 将完整流程封装为 WorkBuddy 技能（名为 tushare-stock-loader），包含 SKILL.md 工作流文档、Python 批处理脚本、交互式 HTML 分析工具和诊断统计清单。在此之后，只需在对话中提及「加载股票数据」或「计算技术指标」等关键词，AI 将自动调用此技能。")

add_heading_styled("4.2 在线部署地址", level=2)
add_para("GitHub 代码仓库（包含全局源码、数据、报告）：")
add_para("https://github.com/837604618-cmyk/AIquant")
add_para("技术指标在线分析工具（上传任意 CSV 即可生成四面板交互图表）：")
add_para("https://ff88feeaf3ed4183b8ad6c361a915199.app.codebuddy.work")
add_para("埃斯顿 K 线看板：")
add_para("https://80d9a8b3e4ff480e9337ff429b683b99.app.codebuddy.work")

add_heading_styled("4.3 AI 协作心得", level=2)
add_para("本次作业全程与 WorkBuddy AI 编程助手协作完成。我的工作方式是：用自然语言描述需求（如「帮我通过 Tushare 获取埃斯顿近一年的交易数据」「请解释 RSI、MACD、布林带的计算方法」），AI 负责执行具体操作——调用数据接口、编写 Python 脚本、生成可视化图表、处理字体和编码问题等。我从中学到了两点：第一，AI 能极大加速从想法到结果的转化，让我可以把精力集中在理解指标原理和解读结果上；第二，AI 生成的结果仍然需要人工验证和调整，比如修复 K 线图的日期空隙、处理中文字体渲染异常、调整报告排版格式等。这种人机协作的模式，可能是未来量化研究的一种常态工作方式。")

# ===== 五 =====
page_break()
add_heading_styled("五、延伸讨论与思考题", level=1)

add_heading_styled("5.1 指标共识与市场自反性", level=2)
add_para("乔治·索罗斯的自反性理论指出，市场参与者的认知与价格之间存在双向反馈作用。当越来越多的机构和个人投资者使用 RSI 的 70/30 阈值、MACD 的金叉/死叉信号、布林带的上下轨触达作为交易决策依据时，这种行为本身是否会改变价格的运动轨迹？例如，当 RSI 接近 70 时，大量交易者可能提前卖出以抢跑信号，导致价格在触及 70 之前就发生反转。这种现象在学术文献中被称为指标拥挤（Indicator Crowding）。")
add_para("思考题 1：在 A 股市场中，技术指标的广泛使用是否已经导致了显著的抢跑效应？是否有实证数据表明某些经典指标的预测效力在近年来出现了下降？")

add_heading_styled("5.2 适应性市场假说 (AMH)", level=2)
add_para("MIT 教授 Andrew Lo 提出的适应性市场假说认为，市场效率并非静态，而是随着参与者、制度和技术的变化而动态调整。技术指标的盈利能力会经历发现 - 利用 - 饱和 - 衰退的循环，但衰退后的指标不会永久失效，当足够多的交易者放弃使用后，其效能可能再度恢复。")
add_para("思考题 2：对于 RSI、MACD、布林带这三个经典指标，它们当前处于 AMH 循环的哪个阶段？是否有学者尝试量化指标在不同市场周期中的有效性变化？")

add_heading_styled("5.3 固定参数的适配性", level=2)
add_para("RSI 的 14 日、MACD 的 (12,26,9)、布林带的 (20, 2) 等参数均产生于 20 世纪后半叶的美国市场环境。在当今 A 股市场中，信息传播速度、交易制度（T+1、涨跌停限制）、投资者结构（散户占比较高）都与美国市场存在显著差异。")
add_para("思考题 3：在 A 股市场中使用这些原厂参数是否仍然合理？是否应该对参数进行本地化调整？这种参数调整是否属于数据挖掘（Data Snooping）的范畴？")

add_heading_styled("5.4 指标组合的信号冗余", level=2)
add_para("本次作业使用了 RSI、MACD、布林带三个指标进行综合分析。然而，这三个指标在数学上并非完全独立：布林带和 MACD 都依赖于移动平均，RSI 和 MACD 都捕捉价格动量。当三者共振发出同向信号时，表面上增加了信心，实际上可能只是同一种信息的不同表达方式。")
add_para("思考题 4：在进行多指标组合分析时，如何评估指标之间的信息冗余程度？是否存在系统性的方法（如主成分分析或信息熵度量）来筛选真正独立的信号源？")

# ===== 参考 =====
page_break()
add_heading_styled("参考资料", level=1)
refs = [
    "[1] Wilder, J. W. (1978). New Concepts in Technical Trading Systems.",
    "[2] Appel, G. (1979). The MACD: A Combo of Indicators.",
    "[3] Bollinger, J. (2001). Bollinger on Bollinger Bands. McGraw-Hill.",
    "[4] Lo, A. W. (2004). The Adaptive Markets Hypothesis. JPM, 30(5), 15-29.",
    "[5] Soros, G. (1987). The Alchemy of Finance. Simon & Schuster.",
    "[6] Murphy, J. J. (1999). Technical Analysis of the Financial Markets.",
    "[7] Tushare Pro - https://tushare.pro/",
    "[8] Apache ECharts - https://echarts.apache.org/",
    "[9] matplotlib - https://matplotlib.org/",
    "[10] pandas - https://pandas.pydata.org/",
]
for r in refs:
    add_para(r)

# 保存
doc.save(OUTPUT)
print(f"[OK] 文档已生成: {OUTPUT}")
